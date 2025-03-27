import docx
# from docx.shared import Pt
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
import os

def extract_logs_from_file(filepath):
    """
    Reads up to 20 lines from the given file.
    Returns:
      parsed_data: list of tokenized lines
      raw_lines:   list of raw lines (strings)
    """
    parsed_data = []
    raw_lines = []
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            lines = file.readlines()
            for line in lines[:20]:  # Extract up to 20 records per file
                raw_lines.append(line.rstrip('\n'))
                parsed_data.append(line.strip().split())
    except PermissionError as e:
        print(f"Permission denied when accessing {filepath}: {e}")
    return parsed_data, raw_lines

def create_table_with_raw_data(doc, category, parsed_data, raw_lines, headers):
    """
    Creates a table that includes all the headers + an extra 'Raw Data' column.
    If headers == 2 (like Field/Value), we do a special Field–Value parse.
    Otherwise, we do the multi-column parse. The last column is always Raw Data.
    """
    doc.add_paragraph(category, style='Heading2')
    
    # We add one extra column for "Raw Data"
    extended_headers = headers + ["Raw Data"]
    table = doc.add_table(rows=1, cols=len(extended_headers))
    table.style = "Table Grid"
    table.autofit = True

    # Fill in the extended header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(extended_headers):
        hdr_cells[i].text = header
        # Make the header text bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # For each line, we place the parsed tokens into columns,
    # then place the raw line in the last column.
    for idx, row_tokens in enumerate(parsed_data):
        # Create a new row in the table
        row_cells = table.add_row().cells
        
        # Grab the raw line for this index
        raw_line = raw_lines[idx] if idx < len(raw_lines) else ""

        # If exactly 2 headers in the original definition, do Field–Value approach
        if len(headers) == 2:
            # If the line is empty, just fill blanks
            if not row_tokens:
                row_cells[0].text = ""
                row_cells[1].text = ""
            else:
                # Last token is Value, preceding tokens are Field
                if len(row_tokens) == 1:
                    field = row_tokens[0]
                    value = ""
                else:
                    field = " ".join(row_tokens[:-1])
                    value = row_tokens[-1]
                row_cells[0].text = field
                row_cells[1].text = value

            # The extra column for raw data
            row_cells[2].text = raw_line
        else:
            # Multi-column approach:
            # If row has fewer items than headers, pad it
            if len(row_tokens) < len(headers):
                row_tokens += [""] * (len(headers) - len(row_tokens))
            # If row has more items, join the extras into the last header column
            elif len(row_tokens) > len(headers):
                row_tokens = row_tokens[:len(headers)-1] + [" ".join(row_tokens[len(headers)-1:])]
            
            # Fill the row cells for each parsed token
            for col_index, item in enumerate(row_tokens):
                row_cells[col_index].text = str(item)
            
            # Fill the last column with raw data
            row_cells[len(headers)].text = raw_line

    doc.add_paragraph('\n')

# Dictionary of log categories and their corresponding filenames
log_files = {
    "Account Information": "watch_logs/account_information.txt",
    "Bluetooth Information": "watch_logs/bluetooth_information.txt",
    "Device Properties": "watch_logs/device_properties.txt",
    "System Logs": "watch_logs/logcat_capture.txt",
    "Sensor Data": "watch_logs/sensor_data.txt"
}

# Expected column headers for each log category
column_headers = {
    "Account Information": ["Field", "Value"],
    "Bluetooth Information": ["Field", "Value"],
    "Device Properties": ["Field", "Value"],
    "System Logs": ["Date", "Time", "PID", "TID", "Log Level", "Component", "Message"],
    "Sensor Data": ["Timestamp", "Sensor Type", "Value"]
}

def main():
    doc = docx.Document()
    doc.add_paragraph("Forensic Log Report", style='Title')

    directory = "./"  # Adjust if needed

    for category, filename in log_files.items():
        filepath = os.path.join(directory, filename)
        if os.path.exists(filepath):
            parsed_data, raw_lines = extract_logs_from_file(filepath)
            headers = column_headers.get(category, ["Data"])
            create_table_with_raw_data(doc, category, parsed_data, raw_lines, headers)
        else:
            doc.add_paragraph(f"No data found for {category}\n", style='Heading2')

    output_filename = "Forensic_Log_Report.docx"
    doc.save(output_filename)
    print(f"Forensic log report generated: {output_filename}")

if __name__ == "__main__":
    main()