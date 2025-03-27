import os
import re
import json
import datetime
import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx.shared import Pt, Inches
from collections import defaultdict

# -----------------------------------------------
# Graph Generation Functions
# -----------------------------------------------
def plot_steps_per_day(filepath, doc):
    """
    Reads a CSV file (skipping the first row), processes it to aggregate total steps per day,
    plots a bar chart, saves the chart as an image, and appends it to the Word document.
    """
    # Read CSV, skipping the first row (metadata)
    df = pd.read_csv(filepath, skiprows=1)
    
    # Rename columns for clarity
    df.columns = ["binning_data", "update_time", "create_time", "source_pkg_name", "source_type",
                  "count", "speed", "distance", "calorie", "deviceuuid", "pkg_name", "datauuid", "day_time"]
    
    # Convert 'datauuid' (milliseconds) to a datetime format
    df["datauuid"] = pd.to_datetime(df["datauuid"], unit="ms", errors='coerce')
    df = df.dropna(subset=["datauuid"])
    
    # Extract the date part for aggregation
    df["date"] = df["datauuid"].dt.date
    
    # Aggregate total steps per day
    daily_steps = df.groupby("date")["count"].sum()
    
    if daily_steps.empty:
        print("No valid data found for plotting steps per day.")
        return
    
    plt.figure(figsize=(12, 6))
    plt.bar(daily_steps.index.astype(str), daily_steps.values, color="skyblue")
    plt.xlabel("Date")
    plt.ylabel("Total Steps")
    plt.title("Steps Per Day Trend")
    plt.xticks(rotation=45)
    plt.grid(axis="y", linestyle="--", alpha=0.7)
    
    steps_image_path = "steps_per_day.png"
    plt.savefig(steps_image_path, bbox_inches='tight')
    plt.close()
    
    # Append the steps chart to the document
    doc.add_paragraph("Steps Per Day Trend", style="Heading2")
    doc.add_picture(steps_image_path, width=Inches(6))
    doc.add_paragraph("\n")

def count_events_by_day_hour(filepath):
    """
    Reads a log file and counts the number of log entries per day and per hour.
    Assumes each log line begins with a timestamp in the format:
        MM-DD HH:MM:SS.xxx
    Returns a dictionary with keys as (day, hour) tuples.
    """
    hour_counts = defaultdict(int)
    pattern = re.compile(r"^(\d{2}-\d{2})\s+(\d{2}):")
    
    with open(filepath, "r", encoding="utf-8", errors="ignore") as file:
        for line in file:
            match = pattern.match(line)
            if match:
                day = match.group(1)    # e.g., "03-08"
                hour = match.group(2)   # e.g., "18"
                key = (day, hour)
                hour_counts[key] += 1
    return hour_counts

def plot_log_events(doc, hour_counts):
    """
    Generates a bar chart for event frequencies (grouped by day and hour),
    saves the chart as an image, and appends it to the Word document.
    """
    sorted_keys = sorted(hour_counts.keys(), key=lambda k: (int(k[0].split('-')[0]),
                                                            int(k[0].split('-')[1]),
                                                            int(k[1])))
    x_labels = [f"{day} {hour}:00" for (day, hour) in sorted_keys]
    y_values = [hour_counts[key] for key in sorted_keys]
    
    plt.figure(figsize=(12, 6))
    plt.bar(x_labels, y_values, color='skyblue')
    plt.xlabel("Time (Day and Hour)")
    plt.ylabel("Number of Events")
    plt.title("Hourly Timeline of Event Frequency")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    
    events_image_path = "event_frequency.png"
    plt.savefig(events_image_path, bbox_inches='tight')
    plt.close()
    
    # Append the event frequency chart to the document
    doc.add_paragraph("Hourly Timeline of Event Frequency", style="Heading2")
    doc.add_picture(events_image_path, width=Inches(6))
    doc.add_paragraph("\n")

# -----------------------------------------------
# Main Function: Append Only the Graphs
# -----------------------------------------------
def main():
    # Since this script is in the "scripts" folder,
    # use "../" to reach files in the project root.
    output_filename = "./Forensic_Log_Report.docx"
    
    # Open the document if it exists; otherwise, create a new one.
    if os.path.exists(output_filename):
        doc = docx.Document(output_filename)
        # Append a page break and header indicating Fgraphs
        doc.add_page_break()
        doc.add_paragraph("Appended Graphs", style="Title")
    else:
        doc = docx.Document()
        doc.add_paragraph("Forensic Log Report - Graphs", style="Title")
    
    # Set base directory to one level up (project root)
    directory = "./"
    
    # Append the log events frequency graph using "logcat_capture.txt"
    log_file = os.path.join(directory, "watch_logs", "logcat_capture.txt")
    if os.path.exists(log_file):
        counts = count_events_by_day_hour(log_file)
        plot_log_events(doc, counts)
    else:
        print(f"Log file not found: {log_file}")
    
    # Append the steps per day graph from the CSV file
    csv_file = os.path.join(directory, "watch_logs", "com.samsung.shealth.step_daily_trend.20250308223308.csv.xls")
    if os.path.exists(csv_file):
        plot_steps_per_day(csv_file, doc)
    else:
        print(f"CSV file not found: {csv_file}")
    
    # Save the updated forensic log report
    doc.save(output_filename)
    print(f"Forensic log report updated: {output_filename}")

if __name__ == "__main__":
    main()
